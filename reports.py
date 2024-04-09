import utils
import beams
import plots
import loadfactors
from PyNite import FEModel3D
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm # to access points, inches, centimeters, and millimeters
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK_TYPE # to access paragraph alignment options
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from io import BytesIO

def beam_report_2D(model: FEModel3D, report_title:str, target_combo:str=None, shear_units=None, shear_direction=None, moment_units=None, moment_direction=None, len_units=None):
    """
    Produces a report including the XXshear/moment/focesXXX
    Takes a non-analyzed FEModel3D
    if target combo = none, max is used
    """

    #create option for shear/moment/axial etc
    
    doc = Document("templates\\empty_PFSE_WB6_report_template.docx")
    
    title = doc.add_paragraph(f"{report_title.title()}", style = "Title")
    intro_heading = doc.add_paragraph("Introduction", style = "Heading 1")
    intro_txt = doc.add_paragraph("The purpose of this report is to provide a brief summary of the shear and moment demands on a 2D beam.", style = "Normal")
    intro_run = intro_txt.add_run("")
    intro_run.add_break(WD_BREAK_TYPE.LINE)
    if target_combo:
        reporting_combo = doc.add_paragraph(f"Reporting Combo: {target_combo}.", style = "Normal")
    else:
        reporting_combo = doc.add_paragraph(f"No target reporting combo chosen, only max diagrams/tables displayed.", style = "Normal")
    reporting_combo_run = reporting_combo.add_run("")
    reporting_combo_run.add_break(WD_BREAK_TYPE.LINE)
    
    table_heading =doc.add_paragraph("Reaction Tables", style = "Heading 1") 
    table_heading_run = table_heading.add_run("")
    table_heading_run.add_break(WD_BREAK_TYPE.LINE)




    model.analyze()
    
    # node_loads = {}
    # for node in model.Nodes:
    #     fy_dict = model.Nodes[node].RxnFY
    #     node_loads.update({f"{node} at target combo ({target_combo})": round(fy_dict[f'{target_combo}']) })
    #     max_combo_val = max(list(fy_dict.values()))
    #     max_combo_name = list(fy_dict.keys())[list(fy_dict.values()).index(max_combo_val)]
    #     node_loads.update({f"{node} at max combo ({max_combo_name})": round(max_combo_val)})
    
    #print(node_loads)
    nodes = list(model.Nodes.keys())
    if target_combo:
        table = doc.add_table(len(nodes)+1, 7)
        table.style = "Grid Table 4"

        table.cell(0,0).add_paragraph("Node#")
        table.cell(0,1).add_paragraph(f"FY @ {target_combo} [{shear_units}]")
        table.cell(0,2).add_paragraph(f"FY @ Max Combo [{shear_units}]")
        table.cell(0,3).add_paragraph("FY Max Combo")
        table.cell(0,4).add_paragraph(f"Mz @ {target_combo} [{moment_units}]")
        table.cell(0,5).add_paragraph(f"Mz @ Max Combo [{moment_units}]")
        table.cell(0,6).add_paragraph("Mz Max Combo")

        for row_num, node in enumerate(nodes, start=1):
            fy_dict = model.Nodes[node].RxnFY
            table.cell(row_num,0).add_paragraph(node)
            table.cell(row_num,1).add_paragraph(f"{round(fy_dict[target_combo])}")
            table.cell(row_num,2).add_paragraph(f"{round(max(list(fy_dict.values())))}")
            table.cell(row_num,3).add_paragraph(f"{list(fy_dict.keys())[list(fy_dict.values()).index(max(list(fy_dict.values())))]}")
            Mz_dict = model.Nodes[node].RxnMZ
            table.cell(row_num,4).add_paragraph(f"{round(Mz_dict[f'{target_combo}'])}")
            table.cell(row_num,5).add_paragraph(f"{round(max(list(Mz_dict.values())))}")
            table.cell(row_num,6).add_paragraph(f"{list(Mz_dict.keys())[list(Mz_dict.values()).index(max(list(Mz_dict.values())))]}")
        else:
            table = doc.add_table(len(nodes)+1, 5)
            table.style = "Grid Table 4"

            table.cell(0,0).add_paragraph("Node#")
            table.cell(0,1).add_paragraph(f"FY @ Max Combo [{shear_units}]")
            table.cell(0,2).add_paragraph("FY Max Combo")
            table.cell(0,3).add_paragraph(f"Mz @ Max Combo [{moment_units}]")
            table.cell(0,4).add_paragraph("Mz Max Combo")

            for row_num, node in enumerate(nodes, start=1):
                fy_dict = model.Nodes[node].RxnFY
                table.cell(row_num,0).add_paragraph(node)
                table.cell(row_num,1).add_paragraph(f"{round(max(list(fy_dict.values())))}")
                table.cell(row_num,2).add_paragraph(f"{list(fy_dict.keys())[list(fy_dict.values()).index(max(list(fy_dict.values())))]}")
                Mz_dict = model.Nodes[node].RxnMZ
                table.cell(row_num,3).add_paragraph(f"{round(max(list(Mz_dict.values())))}")
                table.cell(row_num,4).add_paragraph(f"{list(Mz_dict.keys())[list(Mz_dict.values()).index(max(list(Mz_dict.values())))]}")


    shear_arrays = beams.extract_arrays_all_combos(model, "shear", "Fy")
    moment_arrays = beams.extract_arrays_all_combos(model, "moment", "Mz")
    
    if target_combo:
        
        target_plots = doc.add_paragraph(f"Shear and moment diagrams @ {target_combo}", style = "Heading 1")
        env_shear_x_y = loadfactors.load_combo_array(shear_arrays, target_combo)
        env_moment_x_y = loadfactors.load_combo_array(moment_arrays, target_combo)
        shear_plot = plots.beam_2D_plot(env_shear_x_y, "shear", shear_direction, shear_units, len_units)
        shear_plot_data = BytesIO()
        shear_plot.savefig(shear_plot_data)
        shear_img = doc.add_paragraph()
        shear_img_run = shear_img.add_run()
        shear_img_run.add_picture(shear_plot_data) # Pass the BytesIO to .add_picture
        moment_plot = plots.beam_2D_plot(env_moment_x_y, "moment", moment_direction, moment_units, len_units)
        moment_plot_data = BytesIO()
        moment_plot.savefig(moment_plot_data)
        moment_img = doc.add_paragraph()
        moment_img_run = moment_img.add_run()
        moment_img_run.add_picture(moment_plot_data) # Pass the BytesIO to .add_picture

    target_plots = doc.add_paragraph(f"Shear and moment diagrams env S6 combos", style = "Heading 1")
    env_shear_x_y = loadfactors.envelope_max(shear_arrays)
    env_moment_x_y = loadfactors.envelope_max(moment_arrays)

    # print(f"{env_shear_x_y=}")
    # print(f"{env_moment_x_y=}")

    shear_plot = plots.beam_2D_plot(env_shear_x_y, "shear", shear_direction, shear_units, len_units)
    shear_plot_data = BytesIO()
    shear_plot.savefig(shear_plot_data)
    shear_img = doc.add_paragraph()
    shear_img_run = shear_img.add_run()
    shear_img_run.add_picture(shear_plot_data) # Pass the BytesIO to .add_picture
    moment_plot = plots.beam_2D_plot(env_moment_x_y, "moment", moment_direction, moment_units, len_units)
    moment_plot_data = BytesIO()
    moment_plot.savefig(moment_plot_data)
    moment_img = doc.add_paragraph()
    moment_img_run = moment_img.add_run()
    moment_img_run.add_picture(moment_plot_data) # Pass the BytesIO to .add_picture
    
    
    doc.save("test_report.docx")
    return None

model = beams.load_beam_model("test_data/example_beam_wb6.txt", True)
beam_report_2D(model, "2D beam report", "ULS2", "N", "Fy", "Nmm", "Mz", "mm")

